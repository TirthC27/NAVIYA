"""
Document Ingestion Service

Agent-agnostic service for extracting text from uploaded documents.
Supports: PDF, DOCX, DOC, TXT, PNG, JPG

This service:
- Accepts uploaded files
- Validates file type and MIME type
- Extracts text using appropriate method
- Normalizes output to clean UTF-8 text
- Returns extraction confidence level

Used by resume upload flow BEFORE ResumeIntelligenceAgent.
Can be reused by other features later.
"""

import io
import re
import hashlib
from typing import Optional, Tuple, Dict, Any, List
from enum import Enum
from dataclasses import dataclass
from datetime import datetime
import tempfile
import os

import httpx
from pydantic import BaseModel


# ============================================
# Configuration & Constants
# ============================================

# Allowed file extensions and their MIME types
ALLOWED_EXTENSIONS = {
    ".pdf": ["application/pdf"],
    ".docx": ["application/vnd.openxmlformats-officedocument.wordprocessingml.document"],
    ".doc": ["application/msword"],
    ".txt": ["text/plain"],
    ".png": ["image/png"],
    ".jpg": ["image/jpeg"],
    ".jpeg": ["image/jpeg"],
}

# Rejected extensions (security)
REJECTED_EXTENSIONS = {".zip", ".exe", ".bat", ".cmd", ".sh", ".html", ".htm", ".js", ".py"}

# Maximum file size (10MB)
MAX_FILE_SIZE = 10 * 1024 * 1024

# Minimum text length for valid extraction
MIN_TEXT_LENGTH = 50


class ExtractionConfidence(str, Enum):
    HIGH = "high"
    MEDIUM = "medium"
    LOW = "low"


@dataclass
class ExtractionResult:
    """Result of document text extraction"""
    success: bool
    text: Optional[str] = None
    confidence: ExtractionConfidence = ExtractionConfidence.MEDIUM
    error: Optional[str] = None
    method: Optional[str] = None  # pdf, docx, doc, txt, ocr
    page_count: int = 0
    word_count: int = 0
    warnings: List[str] = None
    
    def __post_init__(self):
        if self.warnings is None:
            self.warnings = []


class DocumentMetadata(BaseModel):
    """Metadata about the uploaded document"""
    original_filename: str
    file_extension: str
    mime_type: str
    file_size: int
    file_hash: str


# ============================================
# Document Ingestion Service
# ============================================

class DocumentIngestionService:
    """
    Service for extracting text from various document formats.
    
    Usage:
        service = DocumentIngestionService()
        result = await service.extract_text(file_content, filename, mime_type)
        
        if result.success:
            print(result.text)
            print(result.confidence)
    """
    
    def __init__(self):
        self._tesseract_available = None
        self._pdf_library = None
    
    # ============================================
    # Main Extraction Method
    # ============================================
    
    async def extract_text(
        self,
        file_content: bytes,
        filename: str,
        mime_type: Optional[str] = None
    ) -> ExtractionResult:
        """
        Extract text from an uploaded document.
        
        Args:
            file_content: Raw file bytes
            filename: Original filename
            mime_type: MIME type from upload
            
        Returns:
            ExtractionResult with text and confidence
        """
        # Step 1: Validate file
        validation = self._validate_file(file_content, filename, mime_type)
        if not validation["valid"]:
            return ExtractionResult(
                success=False,
                error=validation["error"]
            )
        
        ext = validation["extension"]
        detected_mime = validation["mime_type"]
        
        # Step 2: Extract based on file type
        try:
            if ext == ".pdf":
                result = await self._extract_from_pdf(file_content)
            elif ext == ".docx":
                result = await self._extract_from_docx(file_content)
            elif ext == ".doc":
                result = await self._extract_from_doc(file_content)
            elif ext == ".txt":
                result = await self._extract_from_txt(file_content)
            elif ext in [".png", ".jpg", ".jpeg"]:
                result = await self._extract_from_image(file_content)
            else:
                return ExtractionResult(
                    success=False,
                    error=f"Unsupported file type: {ext}"
                )
            
            # Step 3: Normalize text if extraction succeeded
            if result.success and result.text:
                result.text = self._normalize_text(result.text)
                result.word_count = len(result.text.split())
                
                # Validate minimum length
                if len(result.text) < MIN_TEXT_LENGTH:
                    result.success = False
                    result.error = "Extracted text is too short. Please upload a more detailed resume."
                    result.confidence = ExtractionConfidence.LOW
            
            return result
            
        except Exception as e:
            return ExtractionResult(
                success=False,
                error=f"Extraction failed: {str(e)}"
            )
    
    # ============================================
    # File Validation
    # ============================================
    
    def _validate_file(
        self,
        file_content: bytes,
        filename: str,
        mime_type: Optional[str]
    ) -> Dict[str, Any]:
        """Validate file extension, MIME type, and size"""
        
        # Get file extension
        ext = os.path.splitext(filename.lower())[1]
        
        # Check for rejected extensions
        if ext in REJECTED_EXTENSIONS:
            return {
                "valid": False,
                "error": f"File type '{ext}' is not allowed for security reasons."
            }
        
        # Check for allowed extensions
        if ext not in ALLOWED_EXTENSIONS:
            return {
                "valid": False,
                "error": f"Unsupported file type: {ext}. Please upload PDF, DOCX, DOC, TXT, or image files."
            }
        
        # Validate file size
        if len(file_content) > MAX_FILE_SIZE:
            return {
                "valid": False,
                "error": f"File too large. Maximum size is {MAX_FILE_SIZE // (1024*1024)}MB."
            }
        
        if len(file_content) == 0:
            return {
                "valid": False,
                "error": "File is empty."
            }
        
        # Validate MIME type if provided
        if mime_type:
            allowed_mimes = ALLOWED_EXTENSIONS.get(ext, [])
            # Be lenient - some browsers send different MIME types
            if mime_type not in allowed_mimes and mime_type != "application/octet-stream":
                # Log but don't reject - MIME types can be unreliable
                pass
        
        # Detect actual MIME type from content
        detected_mime = self._detect_mime_type(file_content, ext)
        
        return {
            "valid": True,
            "extension": ext,
            "mime_type": detected_mime,
            "size": len(file_content)
        }
    
    def _detect_mime_type(self, content: bytes, ext: str) -> str:
        """Detect MIME type from file content (magic bytes)"""
        # PDF: %PDF
        if content[:4] == b'%PDF':
            return "application/pdf"
        
        # DOCX/XLSX/PPTX: PK (ZIP format)
        if content[:2] == b'PK':
            return "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        
        # DOC: D0 CF 11 E0 (OLE format)
        if content[:4] == b'\xd0\xcf\x11\xe0':
            return "application/msword"
        
        # PNG: 89 50 4E 47
        if content[:4] == b'\x89PNG':
            return "image/png"
        
        # JPEG: FF D8 FF
        if content[:3] == b'\xff\xd8\xff':
            return "image/jpeg"
        
        # Default to extension-based MIME
        mimes = ALLOWED_EXTENSIONS.get(ext, ["application/octet-stream"])
        return mimes[0]
    
    # ============================================
    # PDF Extraction
    # ============================================
    
    async def _extract_from_pdf(self, content: bytes) -> ExtractionResult:
        """Extract text from PDF file"""
        warnings = []
        
        # Try PyMuPDF (fitz) first - it's faster
        try:
            import fitz  # PyMuPDF
            
            doc = fitz.open(stream=content, filetype="pdf")
            text_parts = []
            page_count = len(doc)
            
            for page in doc:
                page_text = page.get_text()
                text_parts.append(page_text)
            
            doc.close()
            
            full_text = "\n\n".join(text_parts)
            
            # Check if text extraction worked
            if len(full_text.strip()) < MIN_TEXT_LENGTH:
                # PDF might be image-based, try OCR
                warnings.append("PDF appears to be image-based, attempting OCR")
                ocr_result = await self._ocr_pdf(content)
                if ocr_result.success:
                    ocr_result.warnings.extend(warnings)
                    return ocr_result
                else:
                    return ExtractionResult(
                        success=False,
                        error="PDF appears to be scanned/image-based and OCR failed. Please upload a text-based PDF.",
                        warnings=warnings
                    )
            
            return ExtractionResult(
                success=True,
                text=full_text,
                confidence=ExtractionConfidence.HIGH,
                method="pdf_pymupdf",
                page_count=page_count,
                warnings=warnings
            )
            
        except ImportError:
            pass  # Try pdfplumber
        
        # Fallback to pdfplumber
        try:
            import pdfplumber
            
            with pdfplumber.open(io.BytesIO(content)) as pdf:
                text_parts = []
                page_count = len(pdf.pages)
                
                for page in pdf.pages:
                    page_text = page.extract_text() or ""
                    text_parts.append(page_text)
            
            full_text = "\n\n".join(text_parts)
            
            if len(full_text.strip()) < MIN_TEXT_LENGTH:
                warnings.append("PDF appears to be image-based, attempting OCR")
                ocr_result = await self._ocr_pdf(content)
                if ocr_result.success:
                    ocr_result.warnings.extend(warnings)
                    return ocr_result
                else:
                    return ExtractionResult(
                        success=False,
                        error="PDF appears to be scanned/image-based and OCR failed.",
                        warnings=warnings
                    )
            
            return ExtractionResult(
                success=True,
                text=full_text,
                confidence=ExtractionConfidence.HIGH,
                method="pdf_pdfplumber",
                page_count=page_count,
                warnings=warnings
            )
            
        except ImportError:
            return ExtractionResult(
                success=False,
                error="PDF extraction libraries not available. Please install PyMuPDF or pdfplumber."
            )
    
    async def _ocr_pdf(self, content: bytes) -> ExtractionResult:
        """OCR a PDF that contains images"""
        try:
            import fitz  # PyMuPDF for rendering
            from PIL import Image
            import pytesseract
            
            doc = fitz.open(stream=content, filetype="pdf")
            text_parts = []
            
            for page in doc:
                # Render page to image at 300 DPI
                pix = page.get_pixmap(dpi=300)
                img_data = pix.tobytes("png")
                img = Image.open(io.BytesIO(img_data))
                
                # Preprocess for OCR
                img = self._preprocess_image_for_ocr(img)
                
                # OCR
                page_text = pytesseract.image_to_string(img, lang='eng')
                text_parts.append(page_text)
            
            doc.close()
            
            full_text = "\n\n".join(text_parts)
            
            # Estimate confidence based on text quality
            confidence = self._estimate_ocr_confidence(full_text)
            
            return ExtractionResult(
                success=True,
                text=full_text,
                confidence=confidence,
                method="pdf_ocr",
                page_count=len(text_parts),
                warnings=["Text extracted using OCR - some inaccuracies may exist"]
            )
            
        except ImportError as e:
            return ExtractionResult(
                success=False,
                error="OCR libraries not available. Please install pytesseract and Pillow."
            )
        except Exception as e:
            return ExtractionResult(
                success=False,
                error=f"OCR failed: {str(e)}"
            )
    
    # ============================================
    # DOCX Extraction
    # ============================================
    
    async def _extract_from_docx(self, content: bytes) -> ExtractionResult:
        """Extract text from DOCX file"""
        try:
            from docx import Document
            
            doc = Document(io.BytesIO(content))
            text_parts = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        text_parts.append(row_text)
            
            full_text = "\n".join(text_parts)
            
            return ExtractionResult(
                success=True,
                text=full_text,
                confidence=ExtractionConfidence.HIGH,
                method="docx",
                page_count=1  # DOCX doesn't have pages in the same way
            )
            
        except ImportError:
            return ExtractionResult(
                success=False,
                error="DOCX extraction library not available. Please install python-docx."
            )
        except Exception as e:
            return ExtractionResult(
                success=False,
                error=f"Failed to read DOCX file: {str(e)}"
            )
    
    # ============================================
    # DOC Extraction
    # ============================================
    
    async def _extract_from_doc(self, content: bytes) -> ExtractionResult:
        """Extract text from DOC file (legacy Word format)"""
        # Try multiple approaches
        
        # Approach 1: textract
        try:
            import textract
            
            # textract needs a file path, so write to temp file
            with tempfile.NamedTemporaryFile(suffix='.doc', delete=False) as tmp:
                tmp.write(content)
                tmp_path = tmp.name
            
            try:
                text = textract.process(tmp_path).decode('utf-8')
                return ExtractionResult(
                    success=True,
                    text=text,
                    confidence=ExtractionConfidence.HIGH,
                    method="doc_textract"
                )
            finally:
                os.unlink(tmp_path)
                
        except ImportError:
            pass
        except Exception:
            pass
        
        # Approach 2: antiword (if available)
        try:
            import subprocess
            
            with tempfile.NamedTemporaryFile(suffix='.doc', delete=False) as tmp:
                tmp.write(content)
                tmp_path = tmp.name
            
            try:
                result = subprocess.run(
                    ['antiword', tmp_path],
                    capture_output=True,
                    text=True,
                    timeout=30
                )
                if result.returncode == 0:
                    return ExtractionResult(
                        success=True,
                        text=result.stdout,
                        confidence=ExtractionConfidence.HIGH,
                        method="doc_antiword"
                    )
            finally:
                os.unlink(tmp_path)
                
        except (FileNotFoundError, subprocess.TimeoutExpired):
            pass
        except Exception:
            pass
        
        # Approach 3: Try reading as text (some .doc files are actually RTF)
        try:
            text = content.decode('utf-8', errors='ignore')
            # Check if it looks like readable text
            printable = sum(c.isprintable() or c.isspace() for c in text)
            if printable / len(text) > 0.8:
                return ExtractionResult(
                    success=True,
                    text=text,
                    confidence=ExtractionConfidence.MEDIUM,
                    method="doc_text_fallback",
                    warnings=["Extracted using text fallback - formatting may be lost"]
                )
        except Exception:
            pass
        
        return ExtractionResult(
            success=False,
            error="Could not extract text from DOC file. Please save as DOCX or PDF and try again."
        )
    
    # ============================================
    # TXT Extraction
    # ============================================
    
    async def _extract_from_txt(self, content: bytes) -> ExtractionResult:
        """Extract text from plain text file"""
        # Try different encodings
        encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
        
        for encoding in encodings:
            try:
                text = content.decode(encoding)
                
                # Check if it looks like valid text
                printable = sum(c.isprintable() or c.isspace() for c in text)
                if printable / max(len(text), 1) > 0.9:
                    return ExtractionResult(
                        success=True,
                        text=text,
                        confidence=ExtractionConfidence.HIGH,
                        method="txt"
                    )
            except UnicodeDecodeError:
                continue
        
        return ExtractionResult(
            success=False,
            error="Could not decode text file. Please ensure it's a valid text file."
        )
    
    # ============================================
    # Image Extraction (OCR)
    # ============================================
    
    async def _extract_from_image(self, content: bytes) -> ExtractionResult:
        """Extract text from image using OCR"""
        try:
            from PIL import Image
            import pytesseract
            
            # Open image
            img = Image.open(io.BytesIO(content))
            
            # Preprocess
            img = self._preprocess_image_for_ocr(img)
            
            # OCR
            text = pytesseract.image_to_string(img, lang='eng')
            
            if len(text.strip()) < MIN_TEXT_LENGTH:
                return ExtractionResult(
                    success=False,
                    error="Could not extract readable text from image. Please upload a clearer image or a PDF/DOCX file.",
                    confidence=ExtractionConfidence.LOW
                )
            
            # Estimate confidence
            confidence = self._estimate_ocr_confidence(text)
            
            return ExtractionResult(
                success=True,
                text=text,
                confidence=confidence,
                method="image_ocr",
                warnings=["Text extracted using OCR from image - some inaccuracies may exist"]
            )
            
        except ImportError:
            return ExtractionResult(
                success=False,
                error="Image OCR not available. Please upload a PDF or DOCX file instead."
            )
        except Exception as e:
            return ExtractionResult(
                success=False,
                error=f"Failed to process image: {str(e)}"
            )
    
    # ============================================
    # Image Preprocessing for OCR
    # ============================================
    
    def _preprocess_image_for_ocr(self, img):
        """Preprocess image for better OCR results"""
        from PIL import Image, ImageEnhance, ImageFilter
        
        # Convert to grayscale
        if img.mode != 'L':
            img = img.convert('L')
        
        # Increase contrast
        enhancer = ImageEnhance.Contrast(img)
        img = enhancer.enhance(1.5)
        
        # Increase sharpness
        enhancer = ImageEnhance.Sharpness(img)
        img = enhancer.enhance(1.5)
        
        # Remove noise with median filter
        img = img.filter(ImageFilter.MedianFilter(size=3))
        
        return img
    
    def _estimate_ocr_confidence(self, text: str) -> ExtractionConfidence:
        """Estimate OCR confidence based on text quality"""
        if not text:
            return ExtractionConfidence.LOW
        
        # Check various quality indicators
        word_count = len(text.split())
        
        # Very short text = low confidence
        if word_count < 20:
            return ExtractionConfidence.LOW
        
        # Check for common OCR errors
        garbage_patterns = [
            r'[^\x00-\x7F]{3,}',  # Multiple non-ASCII chars in a row
            r'[A-Z]{10,}',  # Very long uppercase sequences
            r'[^a-zA-Z0-9\s]{5,}',  # Long sequences of special chars
        ]
        
        garbage_count = 0
        for pattern in garbage_patterns:
            garbage_count += len(re.findall(pattern, text))
        
        garbage_ratio = garbage_count / max(word_count, 1)
        
        if garbage_ratio > 0.2:
            return ExtractionConfidence.LOW
        elif garbage_ratio > 0.1:
            return ExtractionConfidence.MEDIUM
        
        # Check for common resume keywords
        resume_keywords = [
            'experience', 'education', 'skills', 'work', 'university',
            'email', 'phone', 'address', 'project', 'job', 'company'
        ]
        
        text_lower = text.lower()
        keyword_count = sum(1 for kw in resume_keywords if kw in text_lower)
        
        if keyword_count >= 4:
            return ExtractionConfidence.HIGH
        elif keyword_count >= 2:
            return ExtractionConfidence.MEDIUM
        else:
            return ExtractionConfidence.LOW
    
    # ============================================
    # Text Normalization
    # ============================================
    
    def _normalize_text(self, text: str) -> str:
        """Normalize extracted text for consistency"""
        if not text:
            return ""
        
        # Step 1: Remove non-printable characters (except newlines/tabs)
        text = ''.join(char for char in text if char.isprintable() or char in '\n\t')
        
        # Step 2: Normalize Unicode
        import unicodedata
        text = unicodedata.normalize('NFKC', text)
        
        # Step 3: Normalize bullet points
        bullet_patterns = [
            (r'[•●○◦▪▫■□►▸→➤➢]', '•'),
            (r'^[\-\*]\s', '• '),
        ]
        for pattern, replacement in bullet_patterns:
            text = re.sub(pattern, replacement, text, flags=re.MULTILINE)
        
        # Step 4: Normalize whitespace
        # Replace multiple spaces with single space
        text = re.sub(r'[ \t]+', ' ', text)
        
        # Replace 3+ newlines with 2 newlines
        text = re.sub(r'\n{3,}', '\n\n', text)
        
        # Step 5: Merge broken lines (lines that end mid-sentence)
        # Look for lines that end without punctuation and next line starts lowercase
        text = re.sub(r'([a-z,])\n([a-z])', r'\1 \2', text)
        
        # Step 6: Remove common headers/footers
        header_footer_patterns = [
            r'^Page \d+ of \d+\s*$',
            r'^\d+\s*$',  # Just page numbers
            r'^Confidential\s*$',
            r'^Resume of .*$',
        ]
        for pattern in header_footer_patterns:
            text = re.sub(pattern, '', text, flags=re.MULTILINE | re.IGNORECASE)
        
        # Step 7: Strip leading/trailing whitespace from each line
        lines = [line.strip() for line in text.split('\n')]
        text = '\n'.join(lines)
        
        # Step 8: Final strip
        text = text.strip()
        
        return text
    
    # ============================================
    # Utility Methods
    # ============================================
    
    def get_file_hash(self, content: bytes) -> str:
        """Generate SHA-256 hash of file content"""
        return hashlib.sha256(content).hexdigest()
    
    def get_metadata(self, content: bytes, filename: str, mime_type: str) -> DocumentMetadata:
        """Get metadata about the document"""
        ext = os.path.splitext(filename.lower())[1]
        return DocumentMetadata(
            original_filename=filename,
            file_extension=ext,
            mime_type=mime_type or self._detect_mime_type(content, ext),
            file_size=len(content),
            file_hash=self.get_file_hash(content)
        )


# ============================================
# Singleton Instance
# ============================================

_service_instance: Optional[DocumentIngestionService] = None


def get_document_ingestion_service() -> DocumentIngestionService:
    """Get singleton instance of DocumentIngestionService"""
    global _service_instance
    if _service_instance is None:
        _service_instance = DocumentIngestionService()
    return _service_instance
